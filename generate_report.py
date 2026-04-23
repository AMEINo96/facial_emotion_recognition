from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_report():
    doc = Document()

    # --- Title Section ---
    title = doc.add_heading('Facial Emotion Recognition Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    metadata = doc.add_paragraph()
    metadata.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
    metadata.add_run('Hardware: NVIDIA GeForce RTX 3050 Ti Laptop GPU\n')
    metadata.add_run('Framework: PyTorch 2.5.1 + CUDA 12.1\n')
    metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- 1. Project Overview ---
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "This project implements a real-time Facial Emotion Recognition system using Deep Learning. "
        "The goal is to classify human facial expressions into seven distinct emotions: "
        "Angry, Disgust, Fear, Happy, Neutral, Sad, and Surprise."
    )

    # --- 2. Dataset Detail ---
    doc.add_heading('2. Dataset Detail', level=1)
    doc.add_paragraph(
        "The project utilizes the FER-2013 (Facial Expression Recognition 2013) dataset. "
        "The dataset consists of 48x48 pixel grayscale images of faces."
    )
    bullet = doc.add_paragraph(style='List Bullet')
    bullet.add_run('Training Samples: 28,709')
    bullet = doc.add_paragraph(style='List Bullet')
    bullet.add_run('Testing Samples: 7,178')

    # --- 3. Model Architecture ---
    doc.add_heading('3. Model Architecture', level=1)
    doc.add_paragraph(
        "A custom Convolutional Neural Network (CNN) was developed for this task. The architecture includes:"
    )
    arch_list = [
        "3 blocks of Conv2d layers with BatchNormalization and ReLU activation.",
        "MaxPool2d layers for spatial dimension reduction.",
        "Dropout layers (0.25 and 0.5) to prevent overfitting.",
        "Fully connected layers leading to a Softmax output for 7 classes."
    ]
    for item in arch_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    # --- 4. Performance Results ---
    doc.add_heading('4. Performance Results', level=1)
    doc.add_paragraph(
        "The model achieved a final accuracy of 63% on the test set. "
        "Below is the detailed classification report generated during the evaluation phase:"
    )

    # Metrics Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Emotion'
    hdr_cells[1].text = 'Precision'
    hdr_cells[2].text = 'Recall'
    hdr_cells[3].text = 'F1-Score'
    hdr_cells[4].text = 'Support'

    data = [
        ["Angry", "0.56", "0.53", "0.55", "958"],
        ["Disgust", "0.81", "0.23", "0.35", "111"],
        ["Fear", "0.53", "0.27", "0.36", "1024"],
        ["Happy", "0.83", "0.88", "0.85", "1774"],
        ["Neutral", "0.54", "0.67", "0.60", "1233"],
        ["Sad", "0.47", "0.58", "0.52", "1247"],
        ["Surprise", "0.78", "0.71", "0.74", "831"],
        ["", "", "", "", ""],
        ["Overall Accuracy", "", "", "0.63", "7178"]
    ]

    for emotion, p, r, f1, s in data:
        row_cells = table.add_row().cells
        row_cells[0].text = emotion
        row_cells[1].text = p
        row_cells[2].text = r
        row_cells[3].text = f1
        row_cells[4].text = s

    # --- 5. Implementation Records ---
    doc.add_heading('5. Implementation Records', level=1)
    records = doc.add_paragraph()
    records.add_run("GPU Acceleration: ").bold = True
    records.add_run("The training was fully accelerated using the NVIDIA RTX 3050 Ti Laptop GPU via CUDA 12.1.\n")
    records.add_run("Environment: ").bold = True
    records.add_run("Python 3.11 with custom PyTorch wheels installed locally to overcome Python 3.14 compatibility issues.\n")
    records.add_run("Training Regimen: ").bold = True
    records.add_run("Adam optimizer with a learning rate scheduler (ReduceLROnPlateau) and CrossEntropy loss function.")

    # --- 6. Real-Time Application ---
    doc.add_heading('6. Real-Time Application', level=1)
    doc.add_paragraph(
        "A real-time inference application was developed using OpenCV. "
        "The app uses the laptop's webcam to detect faces via Haar Cascades and classifies the "
        "current emotion in real-time at the top of the video feed."
    )

    # Save
    filename = "Facial_Emotion_Recognition_Report.docx"
    doc.save(filename)
    print(f"[+] Report generated successfully as {filename}")

if __name__ == "__main__":
    create_report()
